"""Export a learning session to a .docx study document.

Each answer is written with its verification status, the released answer text, and
every claim with its status and the exact cited quotation + source — so the
export carries the same evidence the app showed.
"""

from __future__ import annotations

import io

from sqlalchemy.ext.asyncio import AsyncSession

from app.models import User
from app.services.audit import load_answer_response
from app.services.sessions import get_owned_session, session_answers

_STATUS_LABEL = {
    "VERIFIED": "Verified",
    "INSUFFICIENT_EVIDENCE": "Insufficient Evidence",
    "CONTRADICTION": "Contradiction Detected",
    "ERROR": "Error",
}


async def export_session_docx(session: AsyncSession, user: User, session_id: str) -> tuple[str, bytes]:
    """Return (filename, docx_bytes). Raises NotFoundError if not owned/found."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    sess = await get_owned_session(session, user, session_id)
    if sess is None:
        from app.api.errors import NotFoundError

        raise NotFoundError("Session not found.")
    answers = await session_answers(session, session_id)

    doc = Document()
    doc.add_heading(sess.title, level=0)
    intro = doc.add_paragraph(
        "Verified learning session. Every released claim below is supported by a cited "
        "quotation from an approved source; the system abstains where evidence is missing."
    )
    intro.runs[0].italic = True

    if not answers:
        doc.add_paragraph("No answers in this session yet.")

    for ans in answers:
        resp = await load_answer_response(session, user, ans.id)
        doc.add_heading(ans.question, level=1)

        status = ans.status.value if hasattr(ans.status, "value") else str(ans.status)
        s = doc.add_paragraph()
        run = s.add_run(f"Status: {_STATUS_LABEL.get(status, status)}")
        run.bold = True
        run.font.size = Pt(11)

        doc.add_paragraph(ans.answer_text or "(no answer)")

        claims = resp.claims if resp else []
        if claims:
            doc.add_heading("Claims & evidence", level=2)
            sources = {s.source_id: s.title for s in (resp.sources if resp else [])}
            for c in claims:
                cstatus = c.status.value if hasattr(c.status, "value") else str(c.status)
                p = doc.add_paragraph(style="List Bullet")
                label = p.add_run(f"[{cstatus}] ")
                label.bold = True
                label.font.color.rgb = RGBColor(0x33, 0x66, 0x33)
                p.add_run(c.text)
                for e in c.evidence:
                    q = doc.add_paragraph(f"“{e.quotation}”  — {sources.get(e.source_id, 'source')}")
                    q.paragraph_format.left_indent = Pt(24)
                    q.runs[0].italic = True

    buf = io.BytesIO()
    doc.save(buf)
    safe_title = "".join(ch if ch.isalnum() or ch in " -_" else "_" for ch in sess.title)[:60]
    return f"{safe_title or 'session'}.docx", buf.getvalue()
